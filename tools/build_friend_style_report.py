from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
TOOLS_DIR = ROOT / "tools"
DOWNLOADS_DIR = ROOT.parent
OUTPUT = DOCS_DIR / "Programming_Certification_Report_Bridget.docx"
COPY_OUTPUT = DOCS_DIR / "Programming_Certification_Report_Bridget_FRIEND_STYLE.docx"

sys.path.insert(0, str(TOOLS_DIR))
import build_visual_certification_report as visual  # noqa: E402


BLUE = "1F4E79"
TEAL = "176B87"
LIGHT_BLUE = "D9EAF7"
PALE = "F6F9FC"
GRAY = "666666"
WHITE = "FFFFFF"

CHATGPT_UML_IMAGES = {
    "uml_use_case": "ChatGPT Image May 9, 2026, 07_33_33 PM.png",
    "uml_class": "ChatGPT Image May 9, 2026, 07_33_28 PM.png",
    "uml_erd": "ChatGPT Image May 9, 2026, 07_33_24 PM.png",
    "uml_sequence_login": "ChatGPT Image May 9, 2026, 07_33_13 PM.png",
    "uml_sequence_borrowing": "ChatGPT Image May 9, 2026, 07_33_08 PM.png",
    "uml_activity_role": "ChatGPT Image May 9, 2026, 07_33_04 PM.png",
}


def copy_user_diagram(key: str, output_name: str) -> Path:
    source = DOWNLOADS_DIR / CHATGPT_UML_IMAGES[key]
    target = visual.ASSET_DIR / output_name
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)
    return target


def draw_user_sidebar(draw, active: str = "Dashboard") -> None:
    visual.rr(draw, (80, 140, 320, 850), "#102A43", "#102A43", 20, 0)
    visual.text(draw, (110, 178), "Universitas XYZ", 24, "#FFFFFF", True)
    menu = ["Dashboard", "Ajukan Peminjaman", "Riwayat Saya"]
    y = 260
    for item in menu:
        fill = "#176B87" if item == active else "#102A43"
        visual.rr(draw, (105, y, 295, y + 54), fill, fill, 12, 0)
        visual.text(draw, (125, y + 16), item, 19, "#D9E7F2")
        y += 78


def draw_user_dashboard_wireframe() -> Path:
    image, draw, path = visual.save_canvas(
        "wireframe_dashboard_user.png",
        "Wireframe - Dashboard Mahasiswa/Dosen",
    )
    draw_user_sidebar(draw, "Dashboard")
    visual.text(draw, (380, 155), "Dashboard MAHASISWA / DOSEN", 32, visual.INK, True)
    visual.text(draw, (380, 200), "User hanya melihat profil dan riwayat peminjaman sendiri", 20, visual.MUTED)

    cards = [
        ("Nama", "Alya Putri"),
        ("NIM/NIK", "220001"),
        ("Total Peminjaman", "2"),
    ]
    x = 380
    for label, value in cards:
        visual.rr(draw, (x, 270, x + 330, 405), visual.PANEL, "#D5DEE8", 18, 2)
        visual.text(draw, (x + 25, 295), label, 20, visual.MUTED)
        visual.text(draw, (x + 25, 335), value, 28, visual.TEAL, True)
        x += 365

    visual.rr(draw, (380, 470, 1460, 810), visual.PANEL, "#D5DEE8", 18, 2)
    visual.text(draw, (410, 505), "Riwayat Peminjaman Saya", 24, visual.BLUE, True)
    headers = ["Ruang", "Tanggal", "Durasi", "Peralatan", "Status"]
    xs = [420, 640, 820, 1000, 1290]
    for x, header in zip(xs, headers):
        visual.text(draw, (x, 565), header, 18, visual.MUTED, True)
    rows = [
        ["Aula Utama", "10 Mei 2026", "2 jam", "-", "MENUNGGU"],
        ["Lab 301", "11 Mei 2026", "2 jam", "Proyektor, Mic", "DISETUJUI"],
    ]
    y = 620
    for row in rows:
        visual.line(draw, (410, y - 18, 1430, y - 18))
        for x, value in zip(xs, row):
            visual.text(draw, (x, y), value, 18, visual.INK)
        y += 72
    image.save(path)
    return path


def draw_borrowing_list_wireframe() -> Path:
    image, draw, path = visual.save_canvas(
        "wireframe_borrowing_list.png",
        "Wireframe - Pencatatan Peminjaman dan Status",
    )
    visual.draw_sidebar(draw, "Peminjaman")
    visual.text(draw, (380, 155), "Pencatatan Peminjaman", 32, visual.INK, True)
    visual.text(draw, (380, 200), "Admin melihat semua transaksi dan mengubah status pengajuan", 20, visual.MUTED)
    visual.rr(draw, (380, 250, 1460, 820), visual.PANEL, "#D5DEE8", 18, 2)
    headers = ["Peminjam", "Ruang", "Jadwal", "Peralatan", "Status", "Update"]
    xs = [420, 610, 790, 980, 1190, 1320]
    for x, header in zip(xs, headers):
        visual.text(draw, (x, 295), header, 17, visual.MUTED, True)
    rows = [
        ["Alya", "Aula", "09:00", "-", "MENUNGGU", "Dropdown"],
        ["Bima", "-", "13:00", "Kamera x1, Mic x2", "DISETUJUI", "Dropdown"],
        ["Alya", "Lab 301", "10:00", "Proyektor x1", "SELESAI", "Waktu kembali"],
    ]
    y = 350
    for row in rows:
        visual.line(draw, (410, y - 18, 1430, y - 18))
        for x, value in zip(xs, row):
            visual.text(draw, (x, y), value, 16, visual.INK)
        y += 86
    visual.rr(draw, (430, 680, 900, 760), "#FFF7ED", "#FED7AA", 14, 2)
    visual.text(draw, (455, 705), "Mahasiswa/dosen tidak melihat kolom Update", 18, "#9A3412", True)
    visual.text(draw, (455, 735), "Riwayat difilter otomatis berdasarkan akun login.", 17, "#9A3412")
    image.save(path)
    return path


def draw_report_wireframe() -> Path:
    image, draw, path = visual.save_canvas("wireframe_report.png", "Wireframe - Laporan dan Export Excel")
    visual.draw_sidebar(draw, "Laporan")
    visual.text(draw, (380, 155), "Laporan Peminjaman", 32, visual.INK, True)
    visual.text(draw, (380, 200), "Admin dapat mengunduh rekap peminjaman menggunakan ExcelJS", 20, visual.MUTED)
    visual.rr(draw, (380, 270, 780, 420), visual.PANEL, "#D5DEE8", 18, 2)
    visual.text(draw, (410, 305), "Total Peminjaman", 21, visual.MUTED)
    visual.text(draw, (410, 350), "3", 42, visual.TEAL, True)
    visual.rr(draw, (820, 270, 1220, 420), visual.PANEL, "#D5DEE8", 18, 2)
    visual.text(draw, (850, 305), "Format Export", 21, visual.MUTED)
    visual.text(draw, (850, 350), ".xlsx", 42, visual.BLUE, True)
    visual.rr(draw, (1260, 290, 1460, 395), visual.TEAL, visual.TEAL, 16, 0)
    visual.text(draw, (1360, 330), "Export", 26, "#FFFFFF", True, "ma")

    visual.rr(draw, (380, 500, 1460, 790), visual.PANEL, "#D5DEE8", 18, 2)
    headers = ["Peminjam", "Jenis Akun", "Ruang", "Peralatan", "Status"]
    xs = [420, 650, 870, 1060, 1300]
    for x, header in zip(xs, headers):
        visual.text(draw, (x, 545), header, 18, visual.MUTED, True)
    rows = [
        ["Alya Putri", "MAHASISWA", "Aula Utama", "-", "MENUNGGU"],
        ["Dr. Bima", "DOSEN", "-", "Kamera + Mic", "DISETUJUI"],
        ["Alya Putri", "MAHASISWA", "Lab 301", "Proyektor + Mic", "SELESAI"],
    ]
    y = 600
    for row in rows:
        visual.line(draw, (410, y - 18, 1430, y - 18))
        for x, value in zip(xs, row):
            visual.text(draw, (x, y), value, 17, visual.INK)
        y += 58
    image.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margin: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, value: str, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True, WHITE, font_size)
        set_cell_shading(table.rows[0].cells[index], BLUE)
        if widths:
            table.rows[0].cells[index].width = Cm(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, False, None, font_size)
            if row_index % 2 == 1:
                set_cell_shading(cells[index], PALE)
            if widths:
                cells[index].width = Cm(widths[index])
    document.add_paragraph()


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)


def add_figure(document: Document, path: Path, caption: str, width: float = 16.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width))
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color in [
        ("Title", 20, BLUE),
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, TEAL),
        ("Heading 3", 11, BLUE),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def section_heading(document: Document, text: str) -> None:
    document.add_heading(text, level=1)


def subsection(document: Document, text: str) -> None:
    document.add_heading(text, level=2)


def build_assets() -> dict[str, Path]:
    return {
        "login": visual.draw_login_wireframe(),
        "admin_dashboard": visual.draw_dashboard_wireframe(),
        "user_dashboard": draw_user_dashboard_wireframe(),
        "borrowing_form": visual.draw_borrowing_wireframe(),
        "crud": visual.draw_crud_wireframe(),
        "borrowing_list": draw_borrowing_list_wireframe(),
        "report": draw_report_wireframe(),
        "erd": visual.draw_erd_diagram(),
        "class": visual.draw_class_diagram(),
        "uml_use_case": copy_user_diagram("uml_use_case", "uml_use_case_chatgpt.png"),
        "uml_class": copy_user_diagram("uml_class", "uml_class_chatgpt.png"),
        "uml_erd": copy_user_diagram("uml_erd", "uml_erd_chatgpt.png"),
        "uml_sequence_login": copy_user_diagram("uml_sequence_login", "uml_sequence_login_chatgpt.png"),
        "uml_sequence_borrowing": copy_user_diagram("uml_sequence_borrowing", "uml_sequence_borrowing_chatgpt.png"),
        "uml_activity_role": copy_user_diagram("uml_activity_role", "uml_activity_role_chatgpt.png"),
    }


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Programming Certification Report Bridget - 0706012219001")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Sistem Peminjaman Ruang dan Peralatan Universitas XYZ")
    subrun.bold = True
    subrun.font.size = Pt(13)
    subrun.font.color.rgb = RGBColor.from_string(TEAL)
    add_table(
        document,
        ["Identitas", "Keterangan"],
        [
            ["Nama Asesi", "Bridget"],
            ["NIM", "0706012219001"],
            ["Skema Sertifikasi", "Okupasi Pemrograman (Programmer)"],
            ["Konteks Unit", "Menggunakan Spesifikasi Program"],
            ["Bentuk Submission", "Proposal, design report, implementation report, dan testing report"],
        ],
        [4.2, 11.8],
        9,
    )


def add_project_description(document: Document) -> None:
    section_heading(document, "I. Deskripsi Proyek")
    add_paragraph(
        document,
        "Proyek ini bertujuan membangun sistem pengelolaan peminjaman ruang dan peralatan "
        "untuk Universitas XYZ. Sebelum aplikasi dibuat, proses peminjaman dilakukan secara "
        "manual sehingga admin sulit memantau jadwal ruang, stok peralatan, status pengajuan, "
        "dan riwayat peminjaman mahasiswa maupun dosen.",
    )
    add_paragraph(
        document,
        "Sistem yang dikembangkan berbasis web dan menggunakan database relasional. Aplikasi "
        "mampu mencatat identitas peminjam, data ruang, data peralatan, transaksi peminjaman, "
        "detail peralatan yang dipinjam, status pengajuan, dan laporan rekap peminjaman.",
    )
    add_paragraph(
        document,
        "Fitur penting pada sistem adalah peminjaman fleksibel. Pengguna dapat meminjam ruang "
        "saja, peralatan saja, atau ruang dan banyak peralatan sekaligus. Sistem juga menyediakan "
        "timetable sederhana agar pengguna dapat melihat waktu peminjaman yang sudah tercatat.",
    )
    add_table(
        document,
        ["Aspek", "Penjelasan"],
        [
            ["Masalah utama", "Pencatatan manual membuat admin sulit mengelola ruang, stok barang, dan status peminjaman."],
            ["Solusi", "Aplikasi web dengan CRUD, validasi input, role access, database relasional, dan export laporan."],
            ["Pengguna", "Admin, mahasiswa, dan dosen."],
            ["Output", "Data master, transaksi peminjaman, timetable, status pengajuan, dan laporan Excel."],
        ],
        [4, 12],
    )
    subsection(document, "Tujuan Sistem")
    add_bullets(
        document,
        [
            "Membantu admin mengelola peminjam, ruang, dan peralatan secara terpusat.",
            "Memastikan pengajuan peminjaman tercatat dengan status yang jelas.",
            "Mencegah input tidak valid sebelum data masuk ke database.",
            "Menyediakan laporan peminjaman yang dapat diekspor menggunakan library pihak ketiga.",
            "Menunjukkan penerapan OOP, relational database, unit testing, dan design specification.",
        ],
    )


def add_tech_stack(document: Document) -> None:
    section_heading(document, "II. Tech Stack")
    add_paragraph(
        document,
        "Proyek ini menggunakan Next.js dengan TypeScript sebagai framework utama. Next.js "
        "dipilih karena dapat menangani halaman, server component, server action, dan route "
        "handler dalam satu project. TypeScript membantu menjaga struktur data lebih jelas "
        "dan cocok untuk dijelaskan pada saat presentasi spesifikasi program.",
    )
    add_table(
        document,
        ["Komponen", "Teknologi", "Alasan Pemilihan"],
        [
            ["Framework", "Next.js 14", "Menyediakan routing, server rendering, server action, dan struktur aplikasi modern."],
            ["Bahasa", "TypeScript", "Memudahkan pendefinisian tipe data dan mengurangi kesalahan penggunaan properti."],
            ["Database", "Supabase PostgreSQL", "Database relasional berbasis cloud yang sesuai dengan kebutuhan ERD."],
            ["ORM", "Prisma", "Model database mudah dibaca, relasi jelas, dan query lebih rapi."],
            ["Validasi", "Zod", "Validasi input dilakukan sebelum service menyimpan data ke database."],
            ["Export", "ExcelJS", "Library pihak ketiga untuk membuat laporan rekap peminjaman dalam format Excel."],
            ["Testing", "Vitest", "Unit test cepat dan cocok untuk fungsi validasi TypeScript."],
            ["UI", "CSS sederhana", "Tampilan ringan, mudah dijelaskan, dan fokus pada fungsi sistem."],
        ],
        [3.2, 4, 8.8],
    )


def add_use_case_section(document: Document, assets: dict[str, Path]) -> None:
    section_heading(document, "III. Use Case Diagram")
    add_paragraph(
        document,
        "Use case diagram digunakan untuk memperlihatkan hubungan antara aktor dan fungsi utama "
        "sistem. Berdasarkan kebutuhan aplikasi, terdapat tiga aktor: Admin, Mahasiswa, dan Dosen. "
        "Admin memiliki akses penuh terhadap master data, status peminjaman, dan laporan, sedangkan "
        "Mahasiswa serta Dosen hanya dapat mengajukan peminjaman dan melihat riwayat sendiri.",
    )
    add_figure(document, assets["uml_use_case"], "Gambar 1. Use case diagram sistem peminjaman ruang dan peralatan.", 16.5)
    add_table(
        document,
        ["Aktor", "Use Case Utama", "Batasan Akses"],
        [
            ["Admin", "Login, kelola peminjam, kelola ruang, kelola peralatan, ajukan peminjaman, update status, catat pengembalian, export laporan.", "Memiliki akses penuh ke fitur manajemen sistem."],
            ["Mahasiswa", "Login, ajukan peminjaman, lihat riwayat peminjaman.", "Hanya dapat memakai profil borrower sendiri dan tidak dapat mengubah status."],
            ["Dosen", "Login, ajukan peminjaman, lihat riwayat peminjaman.", "Hak akses sama seperti mahasiswa untuk fitur peminjaman."],
        ],
        [3, 8.3, 4.7],
        8,
    )


def add_class_diagram_section(document: Document, assets: dict[str, Path]) -> None:
    section_heading(document, "IV. Class Diagram")
    add_paragraph(
        document,
        "Sistem menggunakan pendekatan berlapis. Entitas data dimodelkan pada Prisma schema, "
        "sedangkan logika bisnis ditempatkan pada service class. Dengan pola ini, setiap "
        "bagian memiliki tanggung jawab yang jelas: form mengirim data, action memeriksa akses, "
        "service menjalankan validasi dan query, lalu Prisma berkomunikasi dengan database.",
    )
    add_figure(document, assets["uml_class"], "Gambar 2. Class diagram sistem dan service layer berbasis OOP.", 16.8)
    subsection(document, "A. Model Entitas (Data)")
    add_table(
        document,
        ["Entitas", "Atribut Utama", "Relasi"],
        [
            ["Account", "id, username, passwordHash, role, borrowerId", "Terhubung opsional ke satu Borrower. Digunakan untuk login dan role access."],
            ["Borrower", "id, name, identityNumber, phone, accountType", "Memiliki banyak Borrowing dan dapat memiliki satu Account."],
            ["Room", "id, code, name, capacity, building, floor, status", "Dapat dipakai oleh banyak Borrowing. roomId pada Borrowing bersifat opsional."],
            ["Equipment", "id, code, name, stock, category", "Terhubung ke Borrowing melalui BorrowingEquipment."],
            ["Borrowing", "id, borrowerId, roomId, requestDate, usageDate, durationHours, status, actualReturnTime, purpose", "Dimiliki satu Borrower, opsional memiliki satu Room, dan memiliki banyak BorrowingEquipment."],
            ["BorrowingEquipment", "id, borrowingId, equipmentId, quantity", "Tabel detail untuk many-to-many antara Borrowing dan Equipment."],
        ],
        [3.5, 5.7, 6.8],
        8,
    )
    subsection(document, "B. Komponen Logika Bisnis (Services)")
    add_table(
        document,
        ["Class", "Method", "Fungsi"],
        [
            ["AccountService", "login(input)", "Memvalidasi username dan password, lalu mengembalikan akun yang valid."],
            ["BorrowerService", "findAll(), findById(), create(), update(), delete()", "Mengelola CRUD peminjam mahasiswa dan dosen."],
            ["RoomService", "findAll(), findById(), create(), update(), delete()", "Mengelola CRUD ruang termasuk kapasitas, gedung, lantai, dan status."],
            ["EquipmentService", "findAll(), findById(), create(), update(), delete()", "Mengelola CRUD peralatan, stok, kode, nama, dan kategori."],
            ["BorrowingService", "findAll(), create(), updateStatus()", "Mencatat transaksi peminjaman, validasi stok, dan update status."],
            ["ReportService", "buildBorrowingWorkbook()", "Membuat workbook Excel untuk rekap peminjaman."],
        ],
        [3.5, 5, 7.5],
        8,
    )
    subsection(document, "C. Deskripsi Relasi Antar Kelas")
    add_bullets(
        document,
        [
            "Server action bergantung pada service class untuk memproses data dari form.",
            "Service class bergantung pada schema validasi Zod agar input aman sebelum query database.",
            "Service class menggunakan Prisma Client untuk membaca dan menyimpan data relasional.",
            "BorrowingService berhubungan dengan EquipmentService secara konsep melalui validasi stok peralatan.",
            "ReportService menggunakan BorrowingService agar data laporan selalu berasal dari sumber transaksi yang sama.",
        ],
    )


def add_erd_section(document: Document, assets: dict[str, Path]) -> None:
    section_heading(document, "V. Entity Relationship Diagram")
    add_paragraph(
        document,
        "ERD menggambarkan struktur database relasional yang digunakan pada Supabase PostgreSQL. "
        "Relasi utama terletak pada Borrower, Room, Equipment, Borrowing, dan BorrowingEquipment. "
        "Tabel Account ditambahkan untuk mendukung login dan pembatasan hak akses.",
    )
    add_figure(document, assets["uml_erd"], "Gambar 3. ERD sistem peminjaman ruang dan peralatan.", 16.5)
    add_table(
        document,
        ["Hubungan", "Tipe Relasi", "Penjelasan Fungsional"],
        [
            ["Account -> Borrower", "One-to-zero-or-one", "Admin tidak wajib memiliki profil peminjam, sedangkan mahasiswa/dosen terhubung ke profil Borrower."],
            ["Borrower -> Borrowing", "One-to-many", "Satu peminjam dapat membuat banyak transaksi peminjaman."],
            ["Room -> Borrowing", "One-to-many opsional", "Satu ruang dapat muncul di banyak peminjaman, tetapi transaksi boleh tidak memakai ruang."],
            ["Borrowing -> BorrowingEquipment", "One-to-many", "Satu transaksi dapat memiliki banyak baris detail peralatan."],
            ["Equipment -> BorrowingEquipment", "One-to-many", "Satu jenis peralatan dapat muncul di banyak transaksi berbeda."],
            ["Borrowing <-> Equipment", "Many-to-many", "Relasi many-to-many dipecah oleh tabel BorrowingEquipment yang menyimpan quantity."],
        ],
        [4.2, 4.2, 7.6],
        8,
    )
    subsection(document, "Kesimpulan Relasi")
    add_paragraph(
        document,
        "Relasi many-to-many antara peminjaman dan peralatan wajib dipecah menjadi tabel detail "
        "karena satu transaksi dapat mencakup banyak peralatan, dan satu peralatan dapat dipakai "
        "dalam banyak transaksi. Field roomId pada Borrowing dibuat opsional agar sistem mendukung "
        "peminjaman barang saja tanpa ruang.",
    )


def add_behavior_diagrams(document: Document, assets: dict[str, Path]) -> None:
    section_heading(document, "VI. Sequence Diagram dan Activity Diagram")
    add_paragraph(
        document,
        "Sequence diagram dan activity diagram ditambahkan berdasarkan materi UML. Sequence diagram "
        "menjelaskan interaksi antar objek dari waktu ke waktu, sedangkan activity diagram menjelaskan "
        "alur proses dan percabangan keputusan pada sistem.",
    )
    subsection(document, "A. Sequence Diagram Login")
    add_paragraph(
        document,
        "Sequence login memperlihatkan proses ketika pengguna mengisi username dan password. Data "
        "dikirim ke server action, diperiksa oleh AccountService, dicocokkan dengan password hash, "
        "lalu sistem membuat session cookie jika login valid.",
    )
    add_figure(document, assets["uml_sequence_login"], "Gambar 4. Sequence diagram proses login role-based.", 16.8)
    subsection(document, "B. Sequence Diagram Ajukan Peminjaman")
    add_paragraph(
        document,
        "Sequence ajukan peminjaman memperlihatkan proses dari form hingga database. Alur ini mencakup "
        "penentuan borrower berdasarkan role, validasi Zod, pengecekan stok peralatan, penyimpanan "
        "Borrowing dan BorrowingEquipment, serta skenario error jika data atau stok tidak valid.",
    )
    add_figure(document, assets["uml_sequence_borrowing"], "Gambar 5. Sequence diagram proses ajukan peminjaman.", 16.8)
    subsection(document, "C. Activity Diagram Role-Based Access")
    add_paragraph(
        document,
        "Activity diagram role-based access menunjukkan keputusan sistem setelah membaca session dan "
        "role akun. Admin diarahkan ke dashboard dengan menu penuh, sedangkan mahasiswa dan dosen "
        "mendapat menu terbatas serta tidak dapat mengakses halaman admin.",
    )
    add_figure(document, assets["uml_activity_role"], "Gambar 6. Activity diagram pembatasan akses berdasarkan role.", 16.5)
    add_table(
        document,
        ["Diagram", "Fokus", "Manfaat untuk Laporan"],
        [
            ["Sequence Login", "Interaksi User, Login Page, Server Action, AccountService, Prisma, dan Database.", "Menjelaskan alur autentikasi dan session."],
            ["Sequence Ajukan Peminjaman", "Interaksi Form, Auth, BorrowingService, Zod, Prisma, Supabase, dan validasi stok.", "Menjelaskan alur transaksi utama aplikasi."],
            ["Activity Role Access", "Percabangan akses Admin, Mahasiswa, dan Dosen.", "Membuktikan role access sudah dirancang dengan jelas."],
        ],
        [4, 6.3, 5.7],
        8,
    )


def add_wireframes(document: Document, assets: dict[str, Path]) -> None:
    section_heading(document, "VII. Wireframe")
    add_paragraph(
        document,
        "Wireframe dibuat untuk menjelaskan alur antarmuka utama aplikasi. Fokus rancangan UI "
        "adalah membedakan akses admin dan non-admin, menampilkan informasi role secara jelas, "
        "dan membuat form peminjaman mudah dipahami saat demonstrasi.",
    )
    wireframes = [
        ("login", "A. Login Role-Based", "Gambar 7. Halaman login terpisah dari dashboard dan memakai akun demo admin, mahasiswa, dan dosen."),
        ("admin_dashboard", "B. Dashboard Admin", "Gambar 8. Dashboard admin menampilkan ringkasan data dan jumlah akun mahasiswa/dosen."),
        ("user_dashboard", "C. Dashboard Mahasiswa/Dosen", "Gambar 9. Dashboard user hanya menampilkan profil dan riwayat peminjaman pribadi."),
        ("crud", "D. Manajemen Master Data", "Gambar 10. Halaman CRUD untuk peminjam, ruang, dan peralatan."),
        ("borrowing_form", "E. Form Peminjaman dan Timetable", "Gambar 11. Form dapat digunakan untuk pinjam ruang saja, barang saja, atau keduanya."),
        ("borrowing_list", "F. Pencatatan Peminjaman dan Update Status", "Gambar 12. Admin dapat mengubah status, sedangkan user biasa hanya membaca riwayat sendiri."),
        ("report", "G. Laporan dan Export", "Gambar 13. Halaman laporan menyediakan export Excel menggunakan ExcelJS."),
    ]
    for key, title, caption in wireframes:
        subsection(document, title)
        add_figure(document, assets[key], caption)


def add_manual(document: Document) -> None:
    section_heading(document, "VIII. Manual Penggunaan Aplikasi")
    add_paragraph(
        document,
        "Bagian ini menjelaskan langkah menjalankan aplikasi dan skenario demo yang dapat "
        "ditunjukkan kepada asesor. Manual dibuat singkat agar mudah dipraktikkan saat presentasi.",
    )
    add_numbered(
        document,
        [
            "Buka folder project bnsp pada terminal.",
            "Jalankan npm install jika dependensi belum tersedia.",
            "Siapkan file .env dengan DATABASE_URL Supabase PostgreSQL dan variabel Supabase publik.",
            "Jalankan npm run db:generate untuk membuat Prisma Client.",
            "Jalankan npm run db:push untuk menyinkronkan schema Prisma ke database.",
            "Jalankan npm run db:seed untuk membuat akun demo, ruang, peralatan, dan transaksi simulasi.",
            "Jalankan npm run dev, lalu buka http://localhost:3001 atau port yang aktif.",
        ],
    )
    add_table(
        document,
        ["Akun Demo", "Username", "Password", "Hak Akses"],
        [
            ["Admin", "admin", "admin123", "Mengelola semua data, update status, dan export laporan."],
            ["Mahasiswa", "alya", "mahasiswa123", "Mengajukan peminjaman dan melihat riwayat sendiri."],
            ["Dosen", "bima", "dosen123", "Mengajukan peminjaman dan melihat riwayat sendiri."],
        ],
        [3, 3, 3.2, 6.8],
        8.5,
    )
    subsection(document, "Skenario Demo Singkat")
    add_bullets(
        document,
        [
            "Login sebagai admin, buka dashboard, lalu tunjukkan jumlah peminjam, ruang, peralatan, dan peminjaman.",
            "Buka menu Peminjam, Ruang, dan Peralatan untuk menunjukkan CRUD master data.",
            "Buka menu Peminjaman, lalu tunjukkan kolom update status hanya muncul pada admin.",
            "Login sebagai mahasiswa atau dosen, lalu tunjukkan sidebar lebih terbatas dan riwayat hanya milik akun sendiri.",
            "Buka Ajukan Peminjaman dan demonstrasikan opsi tidak pinjam ruang, dropdown peralatan, dan timetable.",
            "Buka Laporan sebagai admin, lalu export rekap peminjaman ke Excel.",
        ],
    )


def add_implementation(document: Document) -> None:
    section_heading(document, "IX. Implementasi")
    add_paragraph(
        document,
        "Implementasi sistem mengikuti hasil desain. Struktur aplikasi dipisahkan menjadi route "
        "page, form component, server action, service class, validation schema, Prisma model, dan "
        "database. Pemisahan ini membuat kode lebih bersih dan mudah dijelaskan kepada asesor.",
    )
    subsection(document, "A. Prisma Schema dan Database")
    add_paragraph(
        document,
        "File prisma/schema.prisma berisi enum dan model yang menjadi dasar ERD. Enum AccountType "
        "membedakan MAHASISWA dan DOSEN, RoomStatus membedakan status ketersediaan ruang, "
        "BorrowingStatus membedakan status pengajuan, dan UserRole membedakan ADMIN, MAHASISWA, "
        "serta DOSEN. Model Borrowing memiliki roomId opsional supaya transaksi barang saja dapat "
        "disimpan tanpa ruang.",
    )
    subsection(document, "B. Authentication dan Role Access")
    add_paragraph(
        document,
        "Autentikasi dibuat sederhana menggunakan cookie account_id. Fungsi getCurrentAccount membaca "
        "akun aktif, requireAccount memastikan pengguna sudah login, dan requireAdmin memastikan hanya "
        "admin yang dapat mengakses halaman atau action tertentu. Password demo tetap disimpan dalam "
        "bentuk hash, bukan teks asli.",
    )
    subsection(document, "C. Validation Layer")
    add_paragraph(
        document,
        "Validasi input ditempatkan di src/lib/validation.ts menggunakan Zod. Schema memeriksa field "
        "wajib, format tanggal, angka positif untuk durasi dan quantity, stok tidak negatif, status "
        "yang valid, serta aturan bahwa transaksi minimal harus memilih ruang atau peralatan.",
    )
    add_table(
        document,
        ["Validasi", "Aturan", "Alasan"],
        [
            ["Field wajib", "String harus terisi setelah trim", "Mencegah data kosong masuk database."],
            ["Tanggal", "Input dikonversi dan divalidasi sebagai Date", "Mencegah format tanggal salah."],
            ["Durasi", "Harus integer positif", "Durasi nol atau negatif tidak logis."],
            ["Quantity", "Harus integer positif jika barang dipilih", "Mencegah jumlah pinjam kosong atau negatif."],
            ["Stok", "Quantity tidak boleh melebihi stock", "Menjaga ketersediaan peralatan."],
            ["Jenis peminjaman", "Minimal roomId atau equipmentItems", "Mencegah transaksi kosong."],
        ],
        [3.8, 5.8, 6.4],
        8,
    )
    subsection(document, "D. Service Class")
    add_table(
        document,
        ["Service", "Penjelasan Implementasi"],
        [
            ["AccountService", "Menerima username dan password, memvalidasi input, mencari akun melalui Prisma, lalu membandingkan password dengan hash."],
            ["BorrowerService", "Menyediakan findAll, findById, create, update, dan delete untuk data peminjam. Schema borrowerSchema dipakai pada create/update."],
            ["RoomService", "Menyediakan CRUD ruang dan memastikan kapasitas, lantai, gedung, serta status tervalidasi sebelum disimpan."],
            ["EquipmentService", "Menyediakan CRUD peralatan dan memastikan kode, nama, stok, serta kategori valid."],
            ["BorrowingService", "Membuat transaksi peminjaman, mengambil stok peralatan, menjalankan validateEquipmentStock, lalu menyimpan Borrowing dan BorrowingEquipment."],
            ["ReportService", "Membuat workbook ExcelJS dengan kolom Peminjam, Jenis Akun, Ruang, Tanggal Pakai, Durasi, Status, Peralatan, dan Keperluan."],
        ],
        [4, 12],
        8,
    )
    subsection(document, "E. Server Actions")
    add_paragraph(
        document,
        "Server action di src/lib/actions.ts menjadi penghubung antara form dan service. Setiap "
        "action admin seperti createBorrower, updateRoom, deleteEquipment, dan updateBorrowingStatus "
        "memanggil requireAdmin. Action createBorrowing memanggil requireAccount karena mahasiswa dan "
        "dosen tetap boleh membuat pengajuan, tetapi borrowerId dipaksa memakai akun sendiri.",
    )
    subsection(document, "F. Route Page dan View")
    add_table(
        document,
        ["Route", "Fungsi Tampilan", "Akses"],
        [
            ["/login", "Form login username dan password.", "Guest"],
            ["/", "Dashboard admin atau dashboard mahasiswa/dosen sesuai role.", "Login"],
            ["/borrowers", "CRUD peminjam dan jenis akun.", "Admin"],
            ["/rooms", "CRUD ruang, kapasitas, gedung, lantai, dan status.", "Admin"],
            ["/equipment", "CRUD peralatan, stok, dan kategori.", "Admin"],
            ["/borrowings", "Daftar transaksi. Admin melihat semua, user melihat miliknya sendiri.", "Login"],
            ["/borrowings/new", "Form peminjaman ruang, barang, atau keduanya.", "Login"],
            ["/reports", "Ringkasan laporan dan link export Excel.", "Admin"],
            ["/reports/export", "Route handler export workbook Excel.", "Admin"],
        ],
        [3.5, 8.2, 4.3],
        8,
    )
    subsection(document, "G. Seed Data")
    add_paragraph(
        document,
        "Seed data dibuat pada prisma/seed.ts untuk memudahkan demonstrasi. Data awal mencakup akun "
        "admin, akun mahasiswa Alya, akun dosen Dr. Bima, tujuh ruang, sepuluh peralatan, dan tiga "
        "simulasi peminjaman: ruang saja, barang saja, serta ruang dan barang sekaligus.",
    )


def unit_test_rows() -> list[list[str]]:
    return [
        ["UT001", "Required field", "Nama peminjam kosong", "Schema menolak input", "Pass"],
        ["UT002", "Durasi", "durationHours = 0", "Schema menolak durasi nol", "Pass"],
        ["UT003", "Stok peralatan", "Pinjam kamera 3, stok 2", "validateEquipmentStock menolak input", "Pass"],
        ["UT004", "Room-only borrowing", "Ruang dipilih, equipment kosong", "Schema menerima transaksi", "Pass"],
        ["UT005", "Equipment-only borrowing", "roomId kosong, barang dipilih", "Schema menerima transaksi", "Pass"],
        ["UT006", "Empty borrowing", "roomId kosong dan equipment kosong", "Schema menolak transaksi", "Pass"],
        ["UT007", "Valid borrowing", "Ruang dan barang valid", "Schema menerima transaksi", "Pass"],
    ]


def black_box_rows() -> list[list[str]]:
    return [
        ["LU001", "Login", "Admin login valid", "admin/admin123", "Masuk ke dashboard admin", "Pass"],
        ["LU002", "Login", "Mahasiswa login valid", "alya/mahasiswa123", "Masuk ke dashboard mahasiswa", "Pass"],
        ["LU003", "Login", "Dosen login valid", "bima/dosen123", "Masuk ke dashboard dosen", "Pass"],
        ["LU004", "Login", "Password salah", "admin/salah", "Login ditolak", "Pass"],
        ["RA001", "Role Access", "Mahasiswa akses /equipment", "Login alya lalu buka URL", "Diblokir atau diarahkan dari halaman admin", "Pass"],
        ["RA002", "Role Access", "Dosen akses /rooms", "Login bima lalu buka URL", "Diblokir atau diarahkan dari halaman admin", "Pass"],
        ["RA003", "Role Access", "Mahasiswa tidak bisa update status", "Buka /borrowings", "Kolom update status tidak tampil", "Pass"],
        ["BR001", "Peminjam", "Tambah peminjam", "Nama, NIM/NIK, HP, jenis akun valid", "Data tersimpan dan tampil di tabel", "Pass"],
        ["BR002", "Peminjam", "Ubah peminjam", "Edit nomor HP", "Data berhasil diperbarui", "Pass"],
        ["BR003", "Peminjam", "Hapus peminjam", "Klik hapus sebagai admin", "Data terhapus dari daftar", "Pass"],
        ["RM001", "Ruang", "Tambah ruang", "Kode, nama, kapasitas, gedung, lantai valid", "Ruang baru tersimpan", "Pass"],
        ["RM002", "Ruang", "Kapasitas kosong/nol", "capacity = 0", "Validasi menolak input", "Pass"],
        ["RM003", "Ruang", "Ubah status ruang", "TERSEDIA ke TIDAK_TERSEDIA", "Status tersimpan", "Pass"],
        ["EQ001", "Peralatan", "Tambah peralatan", "Kode, nama, stok, kategori valid", "Peralatan tersimpan", "Pass"],
        ["EQ002", "Peralatan", "Stok negatif", "stock = -1", "Validasi menolak input", "Pass"],
        ["EQ003", "Peralatan", "Ubah stok", "stok 5 menjadi 4", "Stok diperbarui", "Pass"],
        ["PM001", "Peminjaman", "Pinjam ruang saja", "Pilih Aula, barang kosong", "Transaksi tersimpan", "Pass"],
        ["PM002", "Peminjaman", "Pinjam barang saja", "Tidak pilih ruang, pilih kamera dan mic", "Transaksi tersimpan", "Pass"],
        ["PM003", "Peminjaman", "Pinjam ruang dan barang", "Pilih Lab 301, proyektor, mic", "Transaksi tersimpan dengan detail barang", "Pass"],
        ["PM004", "Peminjaman", "Tidak pilih ruang dan barang", "Ruang kosong, equipment kosong", "Transaksi ditolak", "Pass"],
        ["PM005", "Peminjaman", "Jumlah barang melebihi stok", "Pinjam kamera lebih dari stok", "Transaksi ditolak", "Pass"],
        ["PM006", "Peminjaman", "Durasi nol", "durationHours = 0", "Transaksi ditolak", "Pass"],
        ["PM007", "Peminjaman", "Timetable tampil", "Buka /borrowings/new", "Jadwal peminjaman tampil dengan mulai dan selesai", "Pass"],
        ["ST001", "Status", "Admin setujui peminjaman", "Status MENUNGGU ke DISETUJUI", "Status berubah", "Pass"],
        ["ST002", "Status", "Admin tolak peminjaman", "Status MENUNGGU ke DITOLAK", "Status berubah", "Pass"],
        ["ST003", "Status", "Admin selesai tanpa waktu kembali", "Status SELESAI, waktu kosong", "Validasi menolak input", "Pass"],
        ["LP001", "Laporan", "Buka halaman laporan", "Login admin buka /reports", "Ringkasan laporan tampil", "Pass"],
        ["LP002", "Laporan", "Export Excel", "Klik export", "File .xlsx terunduh", "Pass"],
        ["LP003", "Laporan", "User akses laporan", "Login alya buka /reports", "Akses diblokir atau diarahkan", "Pass"],
    ]


def add_testing(document: Document) -> None:
    section_heading(document, "X. Test Case")
    subsection(document, "A. Unit Testing")
    add_paragraph(
        document,
        "Unit testing dilakukan menggunakan Vitest. Pengujian berfokus pada validasi input kritis "
        "secara terisolasi, sehingga fungsi validasi dapat dipastikan berjalan tanpa harus membuka UI.",
    )
    add_table(
        document,
        ["ID", "Area", "Input", "Expected Result", "Status"],
        unit_test_rows(),
        [2, 3.2, 4.5, 4.8, 1.5],
        8,
    )
    subsection(document, "B. Black Box Testing")
    add_paragraph(
        document,
        "Black box testing dilakukan dari sudut pandang pengguna akhir. Pengujian mencakup login, "
        "hak akses, CRUD master data, transaksi peminjaman, update status, timetable, dan export laporan.",
    )
    add_table(
        document,
        ["ID", "Modul", "Deskripsi Test Case", "Input / Langkah", "Expected Result", "Status"],
        black_box_rows(),
        [1.7, 2.2, 3.8, 4.2, 4.2, 1.3],
        7.2,
    )


def add_conclusion(document: Document) -> None:
    section_heading(document, "XI. Kesimpulan")
    add_paragraph(
        document,
        "Berdasarkan hasil desain, implementasi, dan pengujian, sistem peminjaman ruang dan "
        "peralatan Universitas XYZ telah memenuhi kebutuhan studi kasus sertifikasi programmer. "
        "Aplikasi sudah menerapkan database relasional, OOP melalui service class, validasi input, "
        "CRUD, transaksi many-to-many, export laporan, role access, dan unit testing.",
    )
    add_table(
        document,
        ["Kebutuhan Sertifikasi", "Bukti Implementasi", "Status"],
        [
            ["Wireframe", "Login, dashboard admin, dashboard user, CRUD, form peminjaman, status, dan laporan.", "Terpenuhi"],
            ["ERD", "Model Account, Borrower, Room, Equipment, Borrowing, dan BorrowingEquipment.", "Terpenuhi"],
            ["Class Diagram", "Service class untuk account, borrower, room, equipment, borrowing, dan report.", "Terpenuhi"],
            ["Tech Stack", "Next.js, TypeScript, Supabase PostgreSQL, Prisma, Zod, ExcelJS, Vitest.", "Terpenuhi"],
            ["CRUD", "Peminjam, ruang, dan peralatan dapat ditambah, dilihat, diubah, dan dihapus.", "Terpenuhi"],
            ["Peminjaman", "Mendukung ruang saja, barang saja, atau keduanya dengan validasi stok.", "Terpenuhi"],
            ["Status", "Admin dapat mengubah status menjadi disetujui, ditolak, atau selesai.", "Terpenuhi"],
            ["Export", "ReportService menggunakan ExcelJS untuk membuat file Excel.", "Terpenuhi"],
            ["Testing", "Vitest unit test dan black box test case disediakan.", "Terpenuhi"],
        ],
        [4.3, 9.2, 2.5],
        8,
    )
    subsection(document, "References")
    add_bullets(
        document,
        [
            "FR.IA.04A - Penjelasan Singkat Proyek.",
            "Lecture note 07 - Software Architecture.",
            "Lecture note 08 - Design Analysis.",
            "Lecture note 10 - Software Implementation.",
            "Lecture note 11 - Software Testing.",
            "Proposal SE (2) sebagai acuan struktur proposal dan requirement.",
            "Source code aplikasi Universitas XYZ Borrowing System.",
        ],
    )


def build_document() -> Path:
    assets = build_assets()
    document = Document()
    configure(document)
    add_title(document)
    add_project_description(document)
    add_tech_stack(document)
    add_use_case_section(document, assets)
    add_class_diagram_section(document, assets)
    add_erd_section(document, assets)
    add_behavior_diagrams(document, assets)
    add_wireframes(document, assets)
    add_manual(document)
    add_implementation(document)
    add_testing(document)
    add_conclusion(document)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    shutil.copyfile(OUTPUT, COPY_OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
