from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Programming_Certification_Report_Bridget.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Certification Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True, color="FFFFFF")
        set_cell_shading(header_cells[index], "1F4E79")
        if widths:
            header_cells[index].width = Cm(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])

    document.add_paragraph()
    return table


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Callout Table"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF3F8")
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_run.font.size = Pt(10)
    paragraph.add_run("\n" + body)
    document.add_paragraph()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F5597"),
        ("Heading 3", 11, "1F4E79"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    if "Certification Table" not in styles:
        table_style = styles.add_style("Certification Table", 3)
        table_style.font.name = "Calibri"
        table_style.font.size = Pt(9)

    if "Callout Table" not in styles:
        callout_style = styles.add_style("Callout Table", 3)
        callout_style.font.name = "Calibri"
        callout_style.font.size = Pt(10)


def add_cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN PENGAJUAN SERTIFIKASI PEMROGRAMAN")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistem Peminjaman Ruang dan Peralatan Universitas XYZ")
    run.bold = True
    run.font.size = Pt(15)

    document.add_paragraph()
    add_table(
        document,
        ["Informasi", "Keterangan"],
        [
            ["Nama Asesi", "Bridget"],
            ["Skema Sertifikasi", "Pemrogram (Programmer)"],
            ["Unit Kompetensi", "J.620100.009.02 - Menggunakan Spesifikasi Program"],
            ["TUK", "Mandiri / Tempat Kerja / Sewaktu"],
            ["Asesor", "[Nama Asesor]"],
            ["Tanggal", "[Tanggal Presentasi]"],
        ],
        [4.5, 11.0],
    )

    statement = document.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.add_run(
        "Dokumen ini disusun sebagai laporan proposal, desain, implementasi, dan pengujian "
        "untuk perangkat lunak yang dikembangkan dalam studi kasus sertifikasi programmer."
    ).italic = True

    document.add_page_break()


def add_executive_summary(document: Document) -> None:
    document.add_heading("1. Ringkasan Eksekutif", level=1)
    document.add_paragraph(
        "Universitas XYZ membutuhkan sistem untuk mengelola peminjaman ruang dan peralatan "
        "yang sebelumnya dilakukan secara manual. Aplikasi yang dibangun adalah sistem web "
        "berbasis Next.js dan TypeScript dengan database relasional Supabase PostgreSQL. "
        "Sistem mendukung manajemen peminjam, ruang, peralatan, transaksi peminjaman, "
        "pembaruan status oleh admin, export laporan Excel, dan unit test validasi input."
    )
    add_callout(
        document,
        "Tujuan Sertifikasi",
        "Laporan ini membuktikan bahwa asesi mampu menggunakan spesifikasi program melalui "
        "penjelasan wireframe, ERD, class diagram, tech stack, implementasi, dan pengujian."
    )


def add_project_overview(document: Document) -> None:
    document.add_heading("2. Proposal Proyek", level=1)
    document.add_heading("2.1 Latar Belakang", level=2)
    document.add_paragraph(
        "Setiap hari mahasiswa dan dosen mengajukan peminjaman ruang kelas, laboratorium, aula, "
        "serta peralatan pendukung seperti proyektor, kamera, mic wireless, tripod, speaker, "
        "dan lighting kit. Proses manual berisiko menimbulkan bentrok jadwal, kesalahan stok, "
        "dan keterlambatan rekap laporan."
    )

    document.add_heading("2.2 Rumusan Masalah", level=2)
    add_bullets(
        document,
        [
            "Bagaimana mencatat peminjaman ruang dan/atau peralatan dalam satu sistem terpadu?",
            "Bagaimana memastikan validasi stok, durasi, tanggal, dan field wajib dilakukan sebelum data masuk database?",
            "Bagaimana membatasi akses agar admin, mahasiswa, dan dosen memiliki hak yang sesuai?",
            "Bagaimana menyediakan laporan peminjaman yang dapat diekspor untuk kebutuhan administrasi?",
        ],
    )

    document.add_heading("2.3 Tujuan", level=2)
    add_bullets(
        document,
        [
            "Membangun aplikasi peminjaman berbasis web dengan database relasional.",
            "Menerapkan OOP melalui service layer untuk setiap domain utama.",
            "Menyediakan CRUD peminjam, ruang, dan peralatan.",
            "Menyediakan transaksi peminjaman ruang saja, peralatan saja, atau keduanya.",
            "Menyediakan export laporan dan unit test validasi kritis.",
        ],
    )

    document.add_heading("2.4 Stakeholder", level=2)
    add_table(
        document,
        ["Stakeholder", "Kebutuhan Utama"],
        [
            ["Admin", "Mengelola master data, memproses status pengajuan, dan export laporan."],
            ["Mahasiswa", "Mengajukan peminjaman dan melihat status pengajuan pribadi."],
            ["Dosen", "Mengajukan peminjaman dan melihat status pengajuan pribadi."],
            ["Asesor", "Menilai spesifikasi program, desain, implementasi, dan testing."],
        ],
        [4.0, 11.5],
    )


def add_requirements(document: Document) -> None:
    document.add_heading("3. Laporan Kebutuhan Sistem", level=1)
    document.add_heading("3.1 Functional Requirements", level=2)
    add_table(
        document,
        ["Kode", "Kebutuhan", "Status"],
        [
            ["FR-01", "CRUD data peminjam mahasiswa/dosen.", "Terpenuhi"],
            ["FR-02", "CRUD data ruang termasuk kapasitas, gedung, lantai, status.", "Terpenuhi"],
            ["FR-03", "CRUD data peralatan termasuk kode, nama, stok, kategori.", "Terpenuhi"],
            ["FR-04", "Transaksi peminjaman dapat berisi ruang saja, barang saja, atau keduanya.", "Terpenuhi"],
            ["FR-05", "Admin dapat mengubah status menjadi disetujui, ditolak, atau selesai.", "Terpenuhi"],
            ["FR-06", "Sistem dapat export laporan peminjaman ke Excel.", "Terpenuhi"],
            ["FR-07", "Unit test menguji validasi input kritis.", "Terpenuhi"],
        ],
        [2.0, 11.0, 2.5],
    )

    document.add_heading("3.2 Non-Functional Requirements", level=2)
    add_table(
        document,
        ["Kategori", "Kebutuhan", "Implementasi"],
        [
            ["Usability", "UI sederhana untuk demo dan input cepat.", "Dashboard, sidebar role-based, form dropdown."],
            ["Security", "Hak akses dibatasi berdasarkan role.", "Server action guard: requireAccount dan requireAdmin."],
            ["Maintainability", "Kode mudah dibaca dan diberi komentar.", "Service layer, validasi terpusat, komentar Indonesia."],
            ["Reliability", "Validasi dilakukan sebelum database.", "Zod schema dan unit test."],
            ["Portability", "Dapat dijalankan sebagai aplikasi web modern.", "Next.js, TypeScript, Supabase PostgreSQL."],
        ],
        [3.0, 6.0, 6.5],
    )

    document.add_heading("3.3 Matriks Hak Akses", level=2)
    add_table(
        document,
        ["Fitur", "Admin", "Mahasiswa", "Dosen"],
        [
            ["Dashboard sistem", "Ya", "Pribadi", "Pribadi"],
            ["CRUD peminjam", "Ya", "Tidak", "Tidak"],
            ["CRUD ruang", "Ya", "Tidak", "Tidak"],
            ["CRUD peralatan", "Ya", "Tidak", "Tidak"],
            ["Ajukan peminjaman", "Ya", "Ya, akun sendiri", "Ya, akun sendiri"],
            ["Update status", "Ya", "Tidak", "Tidak"],
            ["Export laporan", "Ya", "Tidak", "Tidak"],
        ],
        [5.5, 3.0, 3.5, 3.5],
    )

    document.add_heading("3.4 Use Case Summary", level=2)
    add_table(
        document,
        ["Use Case", "Aktor", "Ringkasan Alur"],
        [
            ["Login", "Admin, Mahasiswa, Dosen", "User mengisi username dan password, sistem membuat session."],
            ["Kelola Master Data", "Admin", "Admin menambah, melihat, mengubah, dan menghapus peminjam, ruang, dan peralatan."],
            ["Ajukan Peminjaman", "Admin, Mahasiswa, Dosen", "User memilih ruang dan/atau peralatan, tanggal pakai, durasi, dan keperluan."],
            ["Update Status", "Admin", "Admin mengubah pengajuan menjadi menunggu, disetujui, ditolak, atau selesai."],
            ["Export Laporan", "Admin", "Admin mengunduh rekap peminjaman dalam format Excel."],
        ],
        [4.0, 4.0, 8.0],
    )


def add_design_report(document: Document) -> None:
    document.add_heading("4. Design Report", level=1)
    document.add_heading("4.1 Arsitektur Sistem", level=2)
    document.add_paragraph(
        "Sistem menggunakan layered architecture untuk memisahkan tanggung jawab antarlapisan. "
        "Pemisahan ini mengikuti prinsip separation of concerns agar UI, server action, service, "
        "validasi, dan database dapat dikembangkan lebih mudah."
    )
    add_table(
        document,
        ["Lapisan", "Komponen", "Tanggung Jawab"],
        [
            ["Presentation", "Next.js App Router dan React Components", "Menampilkan dashboard, form CRUD, timetable, dan laporan."],
            ["Action/Controller", "Server Actions dan Route Handler", "Menerima submit form, menjaga role access, redirect, dan export."],
            ["Service Layer", "AccountService, BorrowingService, dll.", "Menjalankan logika OOP untuk CRUD dan transaksi."],
            ["Validation", "Zod Schema", "Memastikan input valid sebelum database."],
            ["Data Access", "Prisma ORM", "Mapping model TypeScript ke PostgreSQL."],
            ["Database", "Supabase PostgreSQL", "Menyimpan data relasional."],
        ],
        [3.0, 5.0, 8.0],
    )

    document.add_heading("4.2 Wireframe Summary", level=2)
    add_table(
        document,
        ["Halaman", "Elemen Utama", "Tujuan"],
        [
            ["Login", "Form username/password dan demo account", "Masuk sesuai role."],
            ["Dashboard Admin", "Kartu statistik, daftar jenis akun", "Melihat ringkasan sistem."],
            ["Dashboard Mahasiswa/Dosen", "Profil dan riwayat pribadi", "Melihat pengajuan sendiri."],
            ["CRUD Master Data", "Tabel, tombol tambah, ubah, hapus", "Mengelola peminjam, ruang, peralatan."],
            ["Form Peminjaman", "Dropdown ruang, dropdown barang, timetable", "Membuat pengajuan peminjaman."],
            ["Laporan", "Tombol export Excel", "Mengunduh rekap peminjaman."],
        ],
        [4.0, 6.0, 6.0],
    )

    document.add_heading("4.3 ERD dan Domain Model", level=2)
    add_table(
        document,
        ["Entity", "Atribut Kunci", "Relasi"],
        [
            ["Account", "username, passwordHash, role", "Opsional terhubung ke Borrower."],
            ["Borrower", "name, identityNumber, phone, accountType", "Memiliki banyak Borrowing."],
            ["Room", "code, name, capacity, building, floor, status", "Opsional dipakai oleh Borrowing."],
            ["Equipment", "code, name, stock, category", "Dipakai oleh BorrowingEquipment."],
            ["Borrowing", "requestDate, usageDate, durationHours, status, purpose", "Menghubungkan Borrower, Room, dan detail Equipment."],
            ["BorrowingEquipment", "quantity", "Tabel penghubung many-to-many Borrowing dan Equipment."],
        ],
        [4.0, 6.0, 6.0],
    )

    document.add_heading("4.4 Class Diagram Summary", level=2)
    add_table(
        document,
        ["Class / Service", "Method Utama", "Peran"],
        [
            ["AccountService", "login()", "Memvalidasi username dan password."],
            ["BorrowerService", "findAll(), findById(), create(), update(), delete()", "CRUD peminjam."],
            ["RoomService", "findAll(), findById(), create(), update(), delete()", "CRUD ruang."],
            ["EquipmentService", "findAll(), findById(), create(), update(), delete()", "CRUD peralatan."],
            ["BorrowingService", "findAll(), create(), updateStatus()", "Transaksi peminjaman dan status."],
            ["ReportService", "buildBorrowingWorkbook()", "Membuat laporan Excel."],
        ],
        [4.0, 6.0, 6.0],
    )

    document.add_heading("4.5 Alur Utama Peminjaman", level=2)
    add_numbered(
        document,
        [
            "User login sebagai admin, mahasiswa, atau dosen.",
            "User membuka form peminjaman.",
            "Sistem menampilkan timetable peminjaman yang sudah tercatat.",
            "User memilih ruang, peralatan, atau keduanya.",
            "Sistem memvalidasi field wajib, tanggal, durasi, dan stok.",
            "Jika valid, sistem menyimpan peminjaman dengan status MENUNGGU.",
            "Admin meninjau pengajuan dan mengubah status sesuai keputusan.",
        ],
    )


def add_implementation_report(document: Document) -> None:
    document.add_heading("5. Implementation Report", level=1)
    document.add_heading("5.1 Tech Stack", level=2)
    add_table(
        document,
        ["Teknologi", "Fungsi"],
        [
            ["Next.js + TypeScript", "Framework full-stack dan type safety."],
            ["Supabase PostgreSQL", "Database relasional online."],
            ["Prisma ORM", "Schema modeling, query, dan relasi database."],
            ["Zod", "Validasi input sebelum menyentuh database."],
            ["ExcelJS", "Library pihak ketiga untuk export laporan Excel."],
            ["Vitest", "Unit testing validasi kritis."],
        ],
        [5.0, 11.0],
    )

    document.add_heading("5.2 Mapping Design to Code", level=2)
    document.add_paragraph(
        "Desain domain dipetakan ke Prisma model dan service class. Entity seperti Borrower, Room, "
        "Equipment, Borrowing, dan BorrowingEquipment direpresentasikan sebagai model database. "
        "Operasi bisnisnya dipisahkan ke service class agar OOP terlihat konsisten dan mudah diuji."
    )

    document.add_heading("5.3 Strategi Validasi", level=2)
    add_bullets(
        document,
        [
            "Field wajib divalidasi dengan requiredString.",
            "Durasi harus angka positif.",
            "Tanggal pengajuan dan tanggal pakai dipaksa menjadi Date yang valid.",
            "Jumlah barang dipastikan tidak melebihi stok tersedia.",
            "Peminjaman boleh ruang saja, barang saja, atau keduanya, tetapi tidak boleh kosong dua-duanya.",
            "Status selesai mewajibkan waktu pengembalian aktual.",
        ],
    )

    document.add_heading("5.4 Clean Code dan Coding Agreement", level=2)
    add_table(
        document,
        ["Aspek", "Kesepakatan"],
        [
            ["Naming", "Class memakai PascalCase; function dan variable memakai camelCase."],
            ["Folder", "app untuk route, components untuk UI, lib untuk helper, services untuk OOP service."],
            ["Function", "Setiap function fokus pada satu tanggung jawab."],
            ["Comments", "Komentar Indonesia digunakan untuk menjelaskan logika penting."],
            ["Validation", "Validasi terpusat di src/lib/validation.ts."],
        ],
        [4.0, 12.0],
    )


def add_testing_report(document: Document) -> None:
    document.add_heading("6. Testing Report", level=1)
    document.add_heading("6.1 Tujuan Testing", level=2)
    document.add_paragraph(
        "Testing dilakukan untuk menunjukkan bahwa aplikasi memenuhi requirement dan untuk menemukan "
        "input yang menghasilkan perilaku salah sebelum sistem digunakan. Pendekatan ini mencakup "
        "validation testing, defect testing, dan unit testing."
    )

    document.add_heading("6.2 Unit Test Summary", level=2)
    add_table(
        document,
        ["Kode", "Skenario", "Input", "Ekspektasi", "Status"],
        [
            ["TC-01", "Field wajib kosong", "name kosong", "Ditolak oleh borrowerSchema", "Pass"],
            ["TC-02", "Durasi nol", "durationHours = 0", "Ditolak oleh borrowingSchema", "Pass"],
            ["TC-03", "Jumlah melebihi stok", "pinjam 3, stok 2", "Ditolak oleh validateEquipmentStock", "Pass"],
            ["TC-04", "Pinjam ruang saja", "roomId ada, equipmentItems kosong", "Diterima", "Pass"],
            ["TC-05", "Pinjam barang saja", "roomId kosong, equipmentItems ada", "Diterima", "Pass"],
            ["TC-06", "Tidak pilih ruang dan barang", "roomId kosong, equipmentItems kosong", "Ditolak", "Pass"],
            ["TC-07", "Data peminjaman valid", "room + equipment + durasi valid", "Diterima", "Pass"],
        ],
        [1.6, 4.2, 4.0, 5.2, 1.5],
    )

    document.add_heading("6.3 Skenario Manual Acceptance Test", level=2)
    add_table(
        document,
        ["Skenario", "Langkah", "Hasil yang Diharapkan"],
        [
            ["Login admin", "Masuk dengan admin/admin123.", "Admin melihat semua menu."],
            ["Login mahasiswa", "Masuk dengan alya/mahasiswa123.", "Mahasiswa hanya melihat dashboard pribadi, ajukan peminjaman, riwayat."],
            ["Login dosen", "Masuk dengan bima/dosen123.", "Dosen hanya melihat dashboard pribadi, ajukan peminjaman, riwayat."],
            ["Update status", "Admin mengubah status peminjaman.", "Status berubah dan non-admin tidak melihat dropdown status."],
            ["Export laporan", "Admin membuka laporan dan klik Export Excel.", "File Excel rekap peminjaman terunduh."],
        ],
        [3.4, 6.0, 6.6],
    )


def add_compliance_and_references(document: Document) -> None:
    document.add_heading("7. Kesimpulan dan Checklist Kepatuhan", level=1)
    add_table(
        document,
        ["Kebutuhan Sertifikasi", "Bukti Implementasi", "Status"],
        [
            ["Wireframe", "Ringkasan halaman login, dashboard, CRUD, peminjaman, laporan.", "Terpenuhi"],
            ["ERD", "Model Account, Borrower, Room, Equipment, Borrowing, BorrowingEquipment.", "Terpenuhi"],
            ["Class Diagram", "Service class untuk domain utama.", "Terpenuhi"],
            ["Tech Stack", "Next.js, TypeScript, Supabase, Prisma, Zod, ExcelJS, Vitest.", "Terpenuhi"],
            ["CRUD", "CRUD peminjam, ruang, dan peralatan.", "Terpenuhi"],
            ["Third-party library", "ExcelJS untuk export laporan dan Zod untuk validasi.", "Terpenuhi"],
            ["Unit Test", "7 test validasi kritis dengan Vitest.", "Terpenuhi"],
        ],
        [5.0, 8.5, 2.5],
    )
    document.add_paragraph(
        "Berdasarkan hasil desain, implementasi, dan pengujian, aplikasi siap dipresentasikan "
        "sebagai bukti penggunaan spesifikasi program pada skema sertifikasi programmer."
    )

    document.add_heading("8. References", level=1)
    add_bullets(
        document,
        [
            "FR.IA.04A. DIT - Daftar Instruksi Terstruktur, Penjelasan Proyek Singkat.",
            "Lecture note 07 - Software Architecture.",
            "Lecture note 08 - Design Analysis.",
            "Lecture note 10 - Software Implementation.",
            "Lecture note 11 - Software Testing.",
            "Proposal SE (2) sebagai acuan struktur proposal.",
            "Dokumen spesifikasi program proyek Universitas XYZ pada folder docs/program-specification.md.",
        ],
    )


def build_document() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_executive_summary(document)
    add_project_overview(document)
    add_requirements(document)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_design_report(document)
    add_implementation_report(document)
    add_testing_report(document)
    add_compliance_and_references(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
