from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "report_assets"
OUTPUT = DOCS_DIR / "Programming_Certification_Report_Bridget_VISUAL.docx"

BLUE = "#1F4E79"
TEAL = "#176B87"
INK = "#172033"
MUTED = "#657085"
LINE = "#CED7E2"
BG = "#F5F7FA"
PANEL = "#FFFFFF"
GREEN = "#1F7A4D"
ORANGE = "#B7791F"
RED = "#B42318"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rr(draw: ImageDraw.ImageDraw, box, fill, outline=LINE, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw: ImageDraw.ImageDraw, xy, value: str, size=28, fill=INK, bold=False, anchor=None):
    draw.text(xy, value, font=font(size, bold), fill=fill, anchor=anchor)


def line(draw: ImageDraw.ImageDraw, xy, fill=LINE, width=3):
    draw.line(xy, fill=fill, width=width)


def save_canvas(name: str, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 960), BG)
    draw = ImageDraw.Draw(image)
    rr(draw, (40, 40, 1560, 920), PANEL, "#E1E7EF", 28, 2)
    text(draw, (80, 78), title, 34, BLUE, True)
    path = ASSET_DIR / name
    return image, draw, path


def draw_sidebar(draw, active="Dashboard"):
    rr(draw, (80, 140, 320, 850), "#102A43", "#102A43", 20, 0)
    text(draw, (110, 178), "Universitas XYZ", 24, "#FFFFFF", True)
    menu = ["Dashboard", "Peminjam", "Ruang", "Peralatan", "Peminjaman", "Laporan"]
    y = 250
    for item in menu:
        fill = "#176B87" if item == active else "#102A43"
        rr(draw, (105, y, 295, y + 54), fill, fill, 12, 0)
        text(draw, (125, y + 16), item, 20, "#D9E7F2")
        y += 68


def draw_login_wireframe():
    image, draw, path = save_canvas("wireframe_login.png", "Wireframe 1 - Login Role-Based")
    rr(draw, (170, 210, 680, 760), "#FFFFFF", "#D5DEE8", 24, 2)
    text(draw, (220, 260), "Login Akun", 34, BLUE, True)
    text(draw, (220, 315), "Masuk sebagai admin, mahasiswa, atau dosen", 20, MUTED)
    text(draw, (220, 390), "Username", 20, MUTED)
    rr(draw, (220, 420, 620, 480), "#F7FAFC", "#CED7E2", 12, 2)
    text(draw, (245, 438), "admin / alya / bima", 20, "#8A96A8")
    text(draw, (220, 520), "Password", 20, MUTED)
    rr(draw, (220, 550, 620, 610), "#F7FAFC", "#CED7E2", 12, 2)
    text(draw, (245, 568), "********", 20, "#8A96A8")
    rr(draw, (220, 650, 620, 715), TEAL, TEAL, 12, 0)
    text(draw, (420, 668), "Login", 24, "#FFFFFF", True, "ma")

    rr(draw, (780, 210, 1390, 760), "#FFFFFF", "#D5DEE8", 24, 2)
    text(draw, (830, 260), "Demo Account", 32, BLUE, True)
    rows = [("ADMIN", "admin", "admin123"), ("MAHASISWA", "alya", "mahasiswa123"), ("DOSEN", "bima", "dosen123")]
    y = 345
    for role, user, password in rows:
        color = GREEN if role == "ADMIN" else TEAL if role == "MAHASISWA" else ORANGE
        rr(draw, (830, y, 1320, y + 86), "#F7FAFC", "#D5DEE8", 14, 2)
        rr(draw, (855, y + 20, 1035, y + 60), color, color, 20, 0)
        text(draw, (945, y + 31), role, 17, "#FFFFFF", True, "ma")
        text(draw, (1070, y + 18), user, 22, INK, True)
        text(draw, (1070, y + 48), password, 18, MUTED)
        y += 115
    image.save(path)
    return path


def draw_dashboard_wireframe():
    image, draw, path = save_canvas("wireframe_dashboard_admin.png", "Wireframe 2 - Dashboard Admin")
    draw_sidebar(draw, "Dashboard")
    text(draw, (380, 160), "Dashboard Admin", 32, INK, True)
    text(draw, (380, 205), "Ringkasan sistem peminjaman ruang dan peralatan", 20, MUTED)
    cards = [("Peminjam", "2"), ("Akun Mahasiswa", "1"), ("Akun Dosen", "1"), ("Ruang", "7"), ("Peralatan", "10"), ("Peminjaman", "3")]
    x, y = 380, 270
    for i, (label, value) in enumerate(cards):
        rr(draw, (x, y, x + 330, y + 135), PANEL, "#D5DEE8", 18, 2)
        text(draw, (x + 25, y + 25), label, 20, MUTED)
        text(draw, (x + 25, y + 63), value, 42, TEAL, True)
        x += 365
        if (i + 1) % 3 == 0:
            x = 380
            y += 170
    rr(draw, (380, 640, 1460, 835), PANEL, "#D5DEE8", 18, 2)
    text(draw, (410, 675), "Jenis Akun Peminjam", 24, BLUE, True)
    for idx, row in enumerate(["Alya Putri     220001     MAHASISWA", "Dr. Bima Santoso     198801012020121001     DOSEN"]):
        yy = 725 + idx * 52
        line(draw, (410, yy, 1430, yy))
        text(draw, (430, yy + 17), row, 21, INK)
    image.save(path)
    return path


def draw_borrowing_wireframe():
    image, draw, path = save_canvas("wireframe_borrowing_form.png", "Wireframe 3 - Form Peminjaman + Timetable")
    draw_sidebar(draw, "Peminjaman")
    text(draw, (380, 155), "Tambah Peminjaman", 32, INK, True)
    text(draw, (380, 200), "Pilih ruang saja, barang saja, atau keduanya", 20, MUTED)

    labels = ["Peminjam", "Ruang", "Tanggal & Jam Pakai", "Durasi Jam"]
    values = ["Alya Putri", "Tidak pinjam ruang / Aula Utama", "10 Mei 2026 09:00", "2"]
    for i, (label, value) in enumerate(zip(labels, values)):
        x = 380 + (i % 2) * 520
        y = 265 + (i // 2) * 110
        text(draw, (x, y), label, 19, MUTED)
        rr(draw, (x, y + 30, x + 460, y + 86), "#F7FAFC", "#CED7E2", 12, 2)
        text(draw, (x + 18, y + 47), value, 18, INK)

    rr(draw, (380, 505, 930, 720), PANEL, "#D5DEE8", 18, 2)
    text(draw, (410, 535), "Dropdown Peralatan", 23, BLUE, True)
    for i in range(3):
        yy = 585 + i * 44
        rr(draw, (410, yy, 720, yy + 34), "#F7FAFC", "#CED7E2", 8, 1)
        text(draw, (425, yy + 8), f"Barang {i+1}: Kamera / Mic / -", 16, INK)
        rr(draw, (740, yy, 880, yy + 34), "#F7FAFC", "#CED7E2", 8, 1)
        text(draw, (755, yy + 8), "Jumlah", 16, MUTED)

    rr(draw, (965, 505, 1460, 720), PANEL, "#D5DEE8", 18, 2)
    text(draw, (995, 535), "Timetable Peminjaman", 23, BLUE, True)
    timetable = ["09:00-11:00  Aula Utama", "13:00-16:00  Kamera + Mic", "10:00-12:00  Lab 301 + Proyektor"]
    for i, item in enumerate(timetable):
        yy = 590 + i * 42
        text(draw, (1000, yy), item, 18, INK)
        rr(draw, (1300, yy - 6, 1415, yy + 28), "#EAF3F8", "#EAF3F8", 16, 0)
        text(draw, (1358, yy + 4), "MENUNGGU", 14, TEAL, True, "ma")

    image.save(path)
    return path


def draw_crud_wireframe():
    image, draw, path = save_canvas("wireframe_crud.png", "Wireframe 4 - CRUD Master Data")
    draw_sidebar(draw, "Peralatan")
    text(draw, (380, 155), "Manajemen Peralatan", 32, INK, True)
    rr(draw, (1240, 145, 1430, 200), TEAL, TEAL, 12, 0)
    text(draw, (1335, 162), "Tambah", 22, "#FFFFFF", True, "ma")
    rr(draw, (380, 250, 1460, 800), PANEL, "#D5DEE8", 18, 2)
    headers = ["Kode", "Nama", "Kategori", "Stok", "Aksi"]
    x_positions = [420, 600, 870, 1110, 1260]
    for x, h in zip(x_positions, headers):
        text(draw, (x, 290), h, 18, MUTED, True)
    rows = [
        ["PRJ-001", "Proyektor Epson", "Presentasi", "5", "Ubah | Hapus"],
        ["CAM-001", "Kamera Sony", "Dokumentasi", "3", "Ubah | Hapus"],
        ["MIC-001", "Mic Wireless", "Audio", "6", "Ubah | Hapus"],
        ["TRP-001", "Tripod Kamera", "Dokumentasi", "5", "Ubah | Hapus"],
    ]
    y = 345
    for row in rows:
        line(draw, (410, y - 16, 1430, y - 16))
        for x, value in zip(x_positions, row):
            text(draw, (x, y), value, 18, INK)
        y += 75
    image.save(path)
    return path


def draw_erd_diagram():
    image, draw, path = save_canvas("diagram_erd.png", "Diagram ERD - Relasi Data Peminjaman")
    boxes = {
        "Account": (120, 190, 430, 350, ["id", "username", "passwordHash", "role", "borrowerId"]),
        "Borrower": (600, 190, 920, 370, ["id", "name", "identityNumber", "phone", "accountType"]),
        "Room": (1080, 190, 1390, 360, ["id", "code", "name", "capacity", "status"]),
        "Borrowing": (600, 520, 930, 760, ["id", "borrowerId", "roomId?", "usageDate", "durationHours", "status", "purpose"]),
        "BorrowingEquipment": (1040, 535, 1430, 720, ["id", "borrowingId", "equipmentId", "quantity"]),
        "Equipment": (105, 535, 430, 735, ["id", "code", "name", "stock", "category"]),
    }
    for name, (x1, y1, x2, y2, attrs) in boxes.items():
        rr(draw, (x1, y1, x2, y2), PANEL, "#AFC1D4", 18, 2)
        rr(draw, (x1, y1, x2, y1 + 48), BLUE, BLUE, 18, 0)
        text(draw, ((x1 + x2) // 2, y1 + 13), name, 19, "#FFFFFF", True, "ma")
        y = y1 + 66
        for attr in attrs:
            text(draw, (x1 + 20, y), attr, 16, INK)
            y += 25
    connectors = [
        ((430, 270), (600, 270), "1 : 0..1"),
        ((760, 370), (760, 520), "1 : *"),
        ((1080, 285), (930, 610), "0..1 : *"),
        ((930, 625), (1040, 625), "1 : *"),
        ((430, 635), (1040, 635), "* : 1"),
    ]
    for start, end, label in connectors:
        line(draw, (*start, *end), "#6B7C8F", 4)
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        rr(draw, (mx - 45, my - 18, mx + 45, my + 18), "#F7FAFC", "#D5DEE8", 12, 1)
        text(draw, (mx, my - 8), label, 14, MUTED, False, "ma")
    image.save(path)
    return path


def draw_class_diagram():
    image, draw, path = save_canvas("diagram_class.png", "Diagram Class - Service Layer OOP")
    classes = [
        ("AccountService", ["login(input)"], 120, 190),
        ("BorrowerService", ["findAll()", "findById()", "create()", "update()", "delete()"], 500, 190),
        ("RoomService", ["findAll()", "findById()", "create()", "update()", "delete()"], 880, 190),
        ("EquipmentService", ["findAll()", "findById()", "create()", "update()", "delete()"], 120, 520),
        ("BorrowingService", ["findAll()", "create()", "updateStatus()"], 500, 520),
        ("ReportService", ["buildBorrowingWorkbook()"], 880, 520),
    ]
    for name, methods, x, y in classes:
        rr(draw, (x, y, x + 310, y + 220), PANEL, "#AFC1D4", 18, 2)
        rr(draw, (x, y, x + 310, y + 52), TEAL, TEAL, 18, 0)
        text(draw, (x + 155, y + 15), name, 18, "#FFFFFF", True, "ma")
        yy = y + 76
        for method in methods:
            text(draw, (x + 25, yy), "+ " + method, 17, INK)
            yy += 28
    line(draw, (655, 520, 655, 410), "#6B7C8F", 4)
    line(draw, (1035, 520, 740, 630), "#6B7C8F", 4)
    text(draw, (690, 455), "service menggunakan Prisma + Zod", 18, MUTED)
    image.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value: str, bold=False, color=None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True, "FFFFFF")
        set_cell_shading(table.rows[0].cells[index], "1F4E79")
        if widths:
            table.rows[0].cells[index].width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
    document.add_paragraph()
    return table


def add_bullets(document: Document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_figure(document: Document, path: Path, caption: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(16.2))
    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def configure(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color in [("Title", 22, "1F4E79"), ("Heading 1", 15, "1F4E79"), ("Heading 2", 12, "2F5597")]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def build_doc():
    assets = [
        draw_login_wireframe(),
        draw_dashboard_wireframe(),
        draw_borrowing_wireframe(),
        draw_crud_wireframe(),
        draw_erd_diagram(),
        draw_class_diagram(),
    ]

    document = Document()
    configure(document)

    for _ in range(2):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN SERTIFIKASI PEMROGRAMAN")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sistem Peminjaman Ruang dan Peralatan Universitas XYZ").bold = True
    add_table(document, ["Field", "Isi"], [
        ["Nama Asesi", "Bridget"],
        ["Skema", "Pemrogram (Programmer)"],
        ["Unit Kompetensi", "J.620100.009.02 - Menggunakan Spesifikasi Program"],
        ["Asesor", "[Nama Asesor]"],
        ["Tanggal", "[Tanggal Presentasi]"],
    ], [4, 12])
    document.add_page_break()

    document.add_heading("1. Ringkasan Proyek", level=1)
    document.add_paragraph(
        "Aplikasi ini menyelesaikan masalah peminjaman ruang dan peralatan di Universitas XYZ. "
        "Sistem mendukung role admin, mahasiswa, dan dosen; CRUD master data; transaksi peminjaman "
        "ruang saja, barang saja, atau keduanya; validasi input; export laporan; dan unit testing."
    )
    add_table(document, ["Komponen", "Implementasi"], [
        ["Frontend/Backend", "Next.js + TypeScript"],
        ["Database", "Supabase PostgreSQL"],
        ["ORM", "Prisma"],
        ["Validasi", "Zod"],
        ["Export", "ExcelJS"],
        ["Testing", "Vitest"],
    ], [5, 11])

    document.add_heading("2. Proposal dan Kebutuhan", level=1)
    add_bullets(document, [
        "Admin mengelola data peminjam, ruang, peralatan, status peminjaman, dan laporan.",
        "Mahasiswa dan dosen dapat mengajukan peminjaman dan melihat riwayat pribadi.",
        "Peminjaman dapat berupa ruang saja, barang saja, atau ruang dan barang sekaligus.",
        "Validasi dilakukan sebelum data disimpan ke database.",
    ])
    add_table(document, ["Fitur", "Admin", "Mahasiswa", "Dosen"], [
        ["CRUD master data", "Ya", "Tidak", "Tidak"],
        ["Ajukan peminjaman", "Ya", "Akun sendiri", "Akun sendiri"],
        ["Update status", "Ya", "Tidak", "Tidak"],
        ["Export laporan", "Ya", "Tidak", "Tidak"],
        ["Lihat riwayat", "Semua", "Sendiri", "Sendiri"],
    ], [5, 3.5, 3.5, 3.5])

    document.add_heading("3. UI Design Wireframe", level=1)
    for asset, caption in [
        (assets[0], "Gambar 1. Wireframe halaman login role-based."),
        (assets[1], "Gambar 2. Wireframe dashboard admin dengan kartu statistik."),
        (assets[2], "Gambar 3. Wireframe form peminjaman dengan dropdown barang dan timetable."),
        (assets[3], "Gambar 4. Wireframe halaman CRUD master data."),
    ]:
        add_figure(document, asset, caption)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("4. Design Report", level=1)
    document.add_paragraph(
        "Desain memakai layered architecture: presentation layer, server action/controller layer, "
        "service layer, validation layer, data access layer, dan database layer. Pola ini dipilih "
        "agar separation of concerns jelas dan sesuai prinsip software architecture."
    )
    add_figure(document, assets[4], "Gambar 5. ERD relasi utama sistem.")
    add_figure(document, assets[5], "Gambar 6. Class diagram service layer berbasis OOP.")

    add_table(document, ["Entity", "Fungsi"], [
        ["Account", "Menyimpan username, password hash, dan role."],
        ["Borrower", "Menyimpan identitas mahasiswa/dosen."],
        ["Room", "Menyimpan data ruang dan ketersediaan."],
        ["Equipment", "Menyimpan stok dan kategori peralatan."],
        ["Borrowing", "Menyimpan transaksi peminjaman."],
        ["BorrowingEquipment", "Detail banyak peralatan dalam satu peminjaman."],
    ], [5, 11])

    document.add_heading("5. Implementation Report", level=1)
    add_table(document, ["Layer", "File/Komponen", "Penjelasan"], [
        ["Presentation", "src/app dan src/components", "Menampilkan halaman dan form."],
        ["Server Action", "src/lib/actions.ts", "Menangani submit form dan role guard."],
        ["Service", "src/services/*.ts", "OOP untuk CRUD dan transaksi."],
        ["Validation", "src/lib/validation.ts", "Validasi Zod sebelum database."],
        ["Database", "prisma/schema.prisma", "Model relasional Supabase PostgreSQL."],
    ], [3.5, 5.5, 7])

    document.add_heading("6. Testing Report", level=1)
    document.add_paragraph(
        "Testing mengikuti konsep validation testing dan defect testing. Unit test dibuat untuk "
        "memastikan validasi kritis berjalan terisolasi sebelum sistem digunakan."
    )
    add_table(document, ["Kode", "Skenario", "Ekspektasi", "Status"], [
        ["TC-01", "Field wajib kosong", "Ditolak", "Pass"],
        ["TC-02", "Durasi nol/negatif", "Ditolak", "Pass"],
        ["TC-03", "Jumlah barang melebihi stok", "Ditolak", "Pass"],
        ["TC-04", "Pinjam ruang saja", "Diterima", "Pass"],
        ["TC-05", "Pinjam barang saja", "Diterima", "Pass"],
        ["TC-06", "Tidak memilih ruang dan barang", "Ditolak", "Pass"],
        ["TC-07", "Data peminjaman valid", "Diterima", "Pass"],
    ], [2, 7, 4, 2.5])

    document.add_heading("7. Checklist Sertifikasi", level=1)
    add_table(document, ["Requirement BNSP", "Bukti", "Status"], [
        ["Wireframe", "4 wireframe UI design", "Terpenuhi"],
        ["ERD", "Diagram ERD visual", "Terpenuhi"],
        ["Class Diagram", "Diagram service OOP", "Terpenuhi"],
        ["Tech Stack", "Next.js, TypeScript, Supabase, Prisma, Zod, ExcelJS, Vitest", "Terpenuhi"],
        ["CRUD", "Peminjam, ruang, peralatan", "Terpenuhi"],
        ["Export", "ExcelJS", "Terpenuhi"],
        ["Unit Test", "7 test Vitest", "Terpenuhi"],
    ], [5, 8.5, 2.5])

    document.add_heading("8. References", level=1)
    add_bullets(document, [
        "FR.IA.04A - Penjelasan Proyek Singkat.",
        "Lecture note 07 - Software Architecture.",
        "Lecture note 08 - Design Analysis.",
        "Lecture note 10 - Software Implementation.",
        "Lecture note 11 - Software Testing.",
        "Proposal SE (2) sebagai acuan struktur proposal.",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
